"""Attachment text extraction — PDF and DOCX."""

import io

import pytest

from app.services.attachment_text import (
    DOCX_MIME,
    extract_attachment_text_sync,
    extract_docx_text,
    normalize_fragmented_text,
)


def test_normalize_fragmented_text_invoice_number():
    raw = "Invoice No: HO-202 512 -01"
    assert normalize_fragmented_text(raw) == "Invoice No: HO-202512-01"


def test_normalize_fragmented_text_month():
    raw = "Date: 1 Dec em ber 2025"
    assert normalize_fragmented_text(raw) == "Date: 1 December 2025"


def test_extract_docx_reimbursement_invoice():
    from docx import Document

    buffer = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Invoice No: HO-202512-01")
    doc.add_paragraph("Date: 1 December 2025")
    doc.add_paragraph("Total: SGD 282.00")
    doc.add_paragraph("From: Marc Michelmann")
    doc.add_paragraph("Category: Home office expense reimbursement")
    doc.save(buffer)

    text = extract_docx_text(buffer.getvalue())
    assert "HO-202512-01" in text
    assert "282.00" in text
    assert "Marc Michelmann" in text
    assert "Home office expense reimbursement" in text


def test_extract_attachment_text_sync_docx():
    from docx import Document

    buffer = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Employee reimbursement total SGD 50.00")
    doc.save(buffer)

    text = extract_attachment_text_sync(content=buffer.getvalue(), mime_type=DOCX_MIME)
    assert text is not None
    assert "50.00" in text
